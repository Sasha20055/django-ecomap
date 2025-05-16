import io
import datetime
from django.contrib import admin
from django.http import HttpResponse
import openpyxl
from docx import Document
from reportlab.pdfgen import canvas

from .models import Location, WasteType, LocationWaste, Review

def export_locations_excel(modeladmin, request, queryset):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(['ID','Name','Address','Lat','Lon','Added By'])
    for loc in queryset:
        ws.append([
            loc.id,
            loc.name,
            loc.address,
            loc.latitude,
            loc.longitude,
            loc.added_by.get_full_name()
        ])
    resp = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    fname = f'locations_{datetime.date.today()}.xlsx'
    resp['Content-Disposition'] = f'attachment; filename="{fname}"'
    wb.save(resp)
    return resp
export_locations_excel.short_description = "Export selected to Excel"

def export_locations_word(modeladmin, request, queryset):
    doc = Document()
    doc.add_heading('Locations', level=1)
    table = doc.add_table(rows=1, cols=6)
    hdr = table.rows[0].cells
    for i, h in enumerate(['ID','Name','Address','Lat','Lon','Added By']):
        hdr[i].text = h
    for loc in queryset:
        row = table.add_row().cells
        row[0].text = str(loc.id)
        row[1].text = loc.name
        row[2].text = loc.address
        row[3].text = str(loc.latitude)
        row[4].text = str(loc.longitude)
        row[5].text = loc.added_by.get_full_name()
    resp = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    fname = f'locations_{datetime.date.today()}.docx'
    resp['Content-Disposition'] = f'attachment; filename="{fname}"'
    doc.save(resp)
    return resp
export_locations_word.short_description = "Export selected to Word"

def export_locations_pdf(modeladmin, request, queryset):
    buf = io.BytesIO()
    p = canvas.Canvas(buf)
    y = 800
    p.setFont('Helvetica-Bold', 14)
    p.drawString(50, y, 'Locations')
    y -= 30
    p.setFont('Helvetica', 12)
    for loc in queryset:
        line = f"{loc.id}. {loc.name} — {loc.address} ({loc.latitude}, {loc.longitude}) by {loc.added_by.get_full_name()}"
        p.drawString(50, y, line[:100])
        y -= 20
        if y < 50:
            p.showPage(); y = 800
    p.save()
    pdf = buf.getvalue()
    buf.close()
    resp = HttpResponse(pdf, content_type='application/pdf')
    fname = f'locations_{datetime.date.today()}.pdf'
    resp['Content-Disposition'] = f'attachment; filename="{fname}"'
    return resp
export_locations_pdf.short_description = "Export selected to PDF"
admin.site.register(WasteType)
admin.site.register(LocationWaste)
admin.site.register(Review)

@admin.register(Location)
class LocationAdmin(admin.ModelAdmin):
    list_display = ('id', 'name', 'address', 'latitude', 'longitude', 'added_by')
    actions = [
        export_locations_excel,  # ваши функции-экшены по экспорту
        export_locations_word,
        export_locations_pdf,
    ]